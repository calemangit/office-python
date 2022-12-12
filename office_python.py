import xlwings as xw
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document

#new_book = xw.Book()

#productos = new_book.sheets.add('productos')

#worksheet = new_book.sheets[0]

#new_book.sheets

#worksheet.range('A1').value = '#'
#worksheet.range('B1').value = 'Producto'

#Lectura y unión de archivos de excel
orders = pd.read_excel('C:\\Users\\Carlos Alemán\\Documents\\dataset.xlsx', sheet_name = 'Orders')
returns = pd.read_excel('C:\\Users\\Carlos Alemán\\Documents\\returns.xlsx', sheet_name = 'Returns')
returned_orders = orders.merge(returns, on ="Order ID")
returned_categories = returned_orders.groupby(['Category'])[['Sales', 'Profit']].apply(sum)

#Creación de gráfico, generación de archivo .png e inserción de gráfico en archivo de excel
fig = plt.figure()
plt.bar(x = returned_categories.index, height = returned_categories['Sales'])
plt.savefig('analysis.png')
book = xw.Book('dataset.xlsx')
book.sheets.add('graph')
target_sheet = book.sheets[0]
target_sheet.pictures.add(fig, name = 'myplot')
#plt.savefig('C:\\Users\\Carlos Alemán\\Documents\\myplot.png')

#Creación de documento en word e inserción de contenido
document = Document
document = Document()
document.add_heading("Top secret report, highly classified", level = 0)
document.add_heading('Execute Summary', level = 1)
document.add_paragraph('The figure below shows the values of returnes sales by category')
document.add_picture('analysis.png')
document.save('amazing_report.doc')

